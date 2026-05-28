import os
import re
from typing import Dict, Any, List

class DocumentParser:
    @staticmethod
    def parse_txt(file_content: bytes) -> Dict[str, Any]:
        """
        Enhanced parser for text/markdown with intelligent section detection.
        Recognizes multiple heading formats and extracts structured content.
        """
        text = file_content.decode("utf-8", errors="ignore")
        
        sections = {}
        current_section = "Introduction"
        current_lines = []
        
        lines = text.split("\n")
        for i, line in enumerate(lines):
            line_strip = line.strip()
            
            # Match multiple heading formats:
            # 1. Markdown: #, ##, ###
            # 2. Numbered: 1., 1.1, 2.3.4
            # 3. ALL CAPS lines
            # 4. Underlined (next line is === or ---)
            
            header_match = None
            
            # Markdown headers
            md_match = re.match(r"^#{1,6}\s+(.+)$", line_strip)
            if md_match:
                header_match = md_match.group(1)
                
            # Numbered sections (1., 1.1, etc.)
            elif re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", line_strip):
                header_match = re.sub(r"^\d+(\.\d+)*\.?\s+", "", line_strip)
                
            # ALL CAPS (at least 3 words, more than 10 chars)
            elif line_strip.isupper() and len(line_strip) > 10 and len(line_strip.split()) >= 3:
                header_match = line_strip.title()
                
            # Underlined headers
            elif i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r"^[=\-]{3,}$", next_line) and line_strip:
                    header_match = line_strip
                    
            if header_match:
                if current_lines:
                    sections[current_section] = "\n".join(current_lines).strip()
                current_section = header_match
                current_lines = []
            else:
                # Skip underline markers
                if not re.match(r"^[=\-]{3,}$", line_strip):
                    current_lines.append(line)
                
        if current_lines:
            sections[current_section] = "\n".join(current_lines).strip()
            
        # Clean empty sections
        sections = {k: v for k, v in sections.items() if v}
        
        if not sections:
            sections["Core Content"] = text
            
        # Extract structured requirements (bullet points, numbered lists)
        requirements = DocumentParser._extract_requirements(text)
        
        return {
            "success": True,
            "text": text,
            "sections": sections,
            "requirements": requirements,
            "char_count": len(text),
            "word_count": len(text.split())
        }
    
    @staticmethod
    def _extract_requirements(text: str) -> List[str]:
        """Extract individual requirements from bullet points and numbered lists"""
        requirements = []
        
        # Match bullet points: -, *, •
        bullet_pattern = r"^[\s]*[-\*•]\s+(.+)$"
        # Match numbered lists: 1., 1), a., a)
        numbered_pattern = r"^[\s]*(?:\d+|[a-z])[\.\)]\s+(.+)$"
        
        for line in text.split("\n"):
            line = line.strip()
            bullet_match = re.match(bullet_pattern, line)
            numbered_match = re.match(numbered_pattern, line)
            
            if bullet_match:
                requirements.append(bullet_match.group(1))
            elif numbered_match:
                requirements.append(numbered_match.group(1))
                
        return requirements

    @staticmethod
    def parse_pdf(file_content: bytes) -> Dict[str, Any]:
        """
        Enhanced PDF parsing with better text extraction using pdfplumber.
        Falls back to PyMuPDF if pdfplumber is unavailable.
        """
        try:
            # Try pdfplumber first (better text extraction)
            import pdfplumber
            import io
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        
            text = "\n\n".join(text_parts)
            return DocumentParser.parse_txt(text.encode("utf-8"))
            
        except ImportError:
            # Fallback to PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_content, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return DocumentParser.parse_txt(text.encode("utf-8"))
            except Exception as e:
                return {
                    "success": False,
                    "error": f"PDF parsing failed: {str(e)}",
                    "text": ""
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF parsing error: {str(e)}",
                "text": ""
            }

    @staticmethod
    def parse_docx(file_content: bytes) -> Dict[str, Any]:
        """
        Enhanced DOCX parsing with support for tables and formatted text.
        """
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                        
            full_text = "\n".join(text_parts)
            return DocumentParser.parse_txt(full_text.encode("utf-8"))
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX parsing error: {str(e)}",
                "text": ""
            }
